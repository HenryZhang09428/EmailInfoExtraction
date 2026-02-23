"""
邮件提取器模块 (Email Extractor Module)
======================================

用于从邮件相关文件中提取结构化信息（`.eml` / `.txt` / `.docx`）。

该文件定位为“轻量编排器”：复杂的解析、清洗、导出等逻辑被拆分到
`core.extractors.email` 子包内的专用类中；本模块负责串联流程、组织 `SourceDoc` 输出，
并保持对外接口（`EmailExtractor.extract()`）不变。

职责划分（按模块）：
- `EmailParser`：解析 MIME、解码邮件头、抽取 text/html 正文与部件元信息
- `AttachmentHandler`：导出附件/内联资源、去重、推断扩展名、文件名清洗
- `ContentCleaner`：HTML→纯文本、线程截断、LLM 输出形状校正、字段值规整
- `DateParser`：日期解析、年份推断、邮件头 Date 解析
- `ResignationDetector`：离职/退工等“离职行”检测（正则 + LLM）
"""

from __future__ import annotations

import json
import os
import re
import uuid
from email import policy
from email.parser import BytesParser
from email.utils import parseaddr
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from core.extractors.base import BaseExtractor
from core.extractors.email.attachment_handler import AttachmentHandler
from core.extractors.email.content_cleaner import ContentCleaner
from core.extractors.email.date_parser import DateParser
from core.extractors.email.email_parser import EmailParser
from core.extractors.email.resignation_detector import ResignationDetector
from core.ir import BlockType, SourceBlock, SourceDoc
from core.llm import LLMClient
from core.logger import get_logger

logger = get_logger(__name__)


# =========================================================================
# 向后兼容：模块级别别名
# =========================================================================
# 历史测试/外部代码可能直接从该模块导入以下符号（此处不展开具体导入语句）。

decode_mime_header = EmailParser.decode_mime_header

_parse_partial_cn_date = DateParser.parse_partial_cn_date
_parse_any_date = DateParser.parse_any_date
_infer_year_from_email_date = DateParser.infer_year_from_email_date
_to_iso_date = DateParser.to_iso_date
_parse_email_header_date = DateParser.parse_email_header_date

is_image_file = AttachmentHandler.is_image_file
_guess_extension = AttachmentHandler.guess_extension
_sanitize_filename = AttachmentHandler.sanitize_filename
_sha256_bytes = AttachmentHandler.sha256_bytes
export_eml_part = AttachmentHandler.export_eml_part
parse_eml = EmailParser.parse


# =========================================================================
# 主抽取器
# =========================================================================

class EmailExtractor(BaseExtractor):
    """
    邮件文件抽取器（`.eml` / `.txt` / `.docx`）。

    提供统一入口 `extract()`，根据扩展名走不同子流程：

    - `.eml`：解析邮件头与正文，导出附件/内联图片（CID），并调用 LLM 抽取正文结构化记录；
      同时可追加“离职行”等特定检测结果，最后组装为 `SourceDoc` 返回。
    - `.txt`：读取纯文本并调用 LLM 抽取结构化 JSON。
    - `.docx`：解析段落/表格/图片占位信息，序列化为 JSON 后调用 LLM 抽取结构化 JSON。

    注意：
    - 本类主要负责编排与组织输出，不在这里堆叠复杂业务细节。
    - `derived_files` 用于记录导出到磁盘的派生文件（例如附件/图片），便于下游处理与追踪。
    """

    def __init__(
        self,
        llm: LLMClient,
        prompts: Optional[dict] = None,
        source_id: Optional[str] = None,
    ) -> None:
        super().__init__(llm, prompts)
        self.source_id = source_id
        self._attachment_handler = AttachmentHandler()
        self._content_cleaner = ContentCleaner()
        self._resignation_detector = ResignationDetector(llm=llm, prompts=prompts)

    # ------------------------------------------------------------------
    # 对外接口
    # ------------------------------------------------------------------

    def extract(self, file_path: str) -> SourceDoc:
        """
        抽取邮件文件内容，并返回统一的 `SourceDoc`。

        参数：
        - file_path：待处理文件路径（支持 `.eml` / `.txt` / `.docx`）。

        返回：
        - `SourceDoc`：包含解析后的 blocks 与最终 extracted JSON。

        异常：
        - `ValueError`：当文件扩展名不在支持范围内时抛出。

        副作用：
        - 每次调用都会清空并重建 `derived_files`（用于记录导出附件/内联资源等派生文件）。
        """
        self.clear_derived_files()

        path = Path(file_path)
        ext = path.suffix.lower()
        filename = path.name

        source_id = self.source_id or str(uuid.uuid4())
        logger.info("Email extraction start: %s (%s)", filename, ext)

        if ext == ".eml":
            source_doc = self._extract_eml(file_path, filename, source_id)
        elif ext == ".txt":
            source_doc = self._extract_txt(file_path, filename, source_id)
        elif ext == ".docx":
            source_doc = self._extract_docx(file_path, filename, source_id)
        else:
            raise ValueError(f"Unsupported email file format: {ext}")

        logger.info(
            "Email extraction end: %s (derived_files=%d)",
            filename,
            len(self.derived_files),
        )
        return source_doc

    # ------------------------------------------------------------------
    # `.eml` 抽取流程
    # ------------------------------------------------------------------

    def _extract_eml(
        self,
        file_path: str,
        filename: str,
        source_id: str,
    ) -> SourceDoc:
        """
        抽取 `.eml` 文件内容。

        该流程会生成多类 `SourceBlock` 块，以便下游既能访问原始内容，也能访问结构化结果：
        - `EMAIL_HEADERS`：邮件头（From/To/Date/Subject 等）
        - `EMAIL_TEXT` / `EMAIL_HTML`：原始 text/html 正文（如存在）
        - `EMAIL_BODY_TEXT`：归一化后的正文纯文本（用于 LLM 抽取与离职检测）
        - `EMAIL_PARTS_SUMMARY`：MIME 部件概览（用于可观测与排障）
        - 导出的附件/内联资源文件块 + CID 引用块
        - `EXTRACTED_JSON`：最终抽取结果（包含告警 `warnings`、上下文与子结果）
        """
        logger.info("EML extract start: %s", filename)

        blocks: List[SourceBlock] = []
        eml_data = EmailParser.parse(file_path)
        logger.info(
            "EML parsed: %s (has_text=%s has_html=%s parts=%d)",
            filename,
            bool(eml_data.get("bodies", {}).get("text")),
            bool(eml_data.get("bodies", {}).get("html")),
            len(eml_data.get("parts") or []),
        )

        # --- 导出附件/邮件部件的缓存目录：按 source_id 隔离，避免不同邮件互相覆盖 ---
        project_root = Path(__file__).parent.parent.parent
        out_dir = str(project_root / ".cache" / "eml_parts" / source_id)

        # --- 基础块：邮件头 + 原始正文（text/html） ---
        order = 1
        blocks.append(
            SourceBlock(order=order, type=BlockType.EMAIL_HEADERS, content=eml_data["headers"], meta={})
        )
        order += 1

        if eml_data["bodies"]["text"]:
            blocks.append(
                SourceBlock(order=order, type=BlockType.EMAIL_TEXT, content=eml_data["bodies"]["text"], meta={})
            )
            order += 1

        if eml_data["bodies"]["html"]:
            blocks.append(
                SourceBlock(order=order, type=BlockType.EMAIL_HTML, content=eml_data["bodies"]["html"], meta={})
            )
            order += 1

        # --- 归一化正文：生成可用于抽取/检测的正文纯文本，并携带元信息（如截断情况） ---
        normalized_body_text, body_text_meta = ContentCleaner.normalize_body_text(eml_data, max_chars=15000)
        blocks.append(
            SourceBlock(order=order, type=BlockType.EMAIL_BODY_TEXT, content=normalized_body_text, meta=body_text_meta)
        )
        order += 1

        # --- 部件概览：只保留关键字段，便于调试与追溯 ---
        parts_summary = [
            {
                "content_type": p.get("content_type", ""),
                "filename": p.get("filename", ""),
                "content_disposition": p.get("content_disposition", ""),
                "content_id": p.get("content_id", ""),
                "size_bytes": p.get("size_bytes", 0),
            }
            for p in eml_data["parts"]
        ]
        blocks.append(
            SourceBlock(order=order, type=BlockType.EMAIL_PARTS_SUMMARY, content=parts_summary, meta={})
        )
        order += 1

        # --- 从 HTML 中提取 `cid:` 引用（通常是内联图片），用于后续做 CID→文件路径映射 ---
        cid_list: List[str] = []
        if eml_data["bodies"]["html"]:
            cid_pattern = r'<img[^>]+src=["\']cid:([^"\']+)["\']'
            for m in re.finditer(cid_pattern, eml_data["bodies"]["html"], re.IGNORECASE):
                cid = m.group(1)
                if cid not in cid_list:
                    cid_list.append(cid)

        # --- 导出附件/内联资源：遍历 MIME 部件写入磁盘，同时追加对应文件块 ---
        order = self._export_attachments(file_path, out_dir, blocks, order, cid_list)

        # --- CID 引用块：记录 HTML 中引用的 CID，并在可能时回填对应文件路径 ---
        cid_to_path = self._last_cid_to_path  # 由 _export_attachments() 填充
        inline_images: List[dict] = []
        for cid in cid_list:
            blocks.append(
                SourceBlock(order=order, type=BlockType.INLINE_IMAGE_REF, content={"cid": cid}, meta={})
            )
            order += 1

            matched_path = self._resolve_cid(cid, cid_to_path)
            if matched_path:
                blocks.append(
                    SourceBlock(
                        order=order,
                        type=BlockType.INLINE_IMAGE_FILE,
                        content={"cid": cid, "path": matched_path},
                        meta={},
                    )
                )
                order += 1
                inline_images.append({"cid": cid, "path": matched_path})

        # --- 告警汇总：CID 引用缺失、导出失败等诊断信息（不影响主流程返回） ---
        warnings: List[str] = []
        for cid in cid_list:
            if not self._resolve_cid(cid, cid_to_path):
                warnings.append(f"CID {cid} referenced in HTML but not found in email parts")
        if self._last_export_warnings:
            warnings.extend(self._last_export_warnings)

        # --- 邮件上下文摘要：供最终抽取结果 JSON 挂载（避免下游重复解析） ---
        sender_name, sender_email = parseaddr(eml_data["headers"]["from"])
        recipients = []
        for addr in eml_data["headers"]["to"].split(","):
            name, email_addr = parseaddr(addr.strip())
            if email_addr:
                recipients.append({"name": name, "email": email_addr})

        bodies_summary: Dict[str, int] = {}
        if eml_data["bodies"]["text"]:
            bodies_summary["text_length"] = len(eml_data["bodies"]["text"])
        if eml_data["bodies"]["html"]:
            bodies_summary["html_length"] = len(eml_data["bodies"]["html"])

        email_content = {
            "headers": eml_data["headers"],
            "bodies_summary": bodies_summary,
            "inline_images": inline_images,
            "attachments_summary": self._last_attachments_summary,
            "warnings": warnings,
        }

        # === 离职/退工等“离职行”检测：基于正文 + 邮件头 Date（用于时间上下文/年份推断） ===
        email_header_date = eml_data.get("headers", {}).get("date", "")
        email_dt = DateParser.parse_email_header_date(email_header_date)
        leave_lines_extracted = self._resignation_detector.run(normalized_body_text, email_dt, filename)

        # === 正文通用信息抽取：使用 LLM 将正文转换为结构化记录（与具体模板解耦） ===
        body_extracted = self._llm_extract_body(normalized_body_text, body_text_meta, filename)

        # --- 组装最终抽取结果 JSON：合并告警、上下文信息与子流程结果 ---
        top_warnings: List[str] = []
        for w in (email_content.get("warnings") or []) + (body_extracted.get("warnings") or []):
            if isinstance(w, str) and w and w not in top_warnings:
                top_warnings.append(w)

        extracted_json: Dict[str, Any] = {
            "data": body_extracted.get("data", []),
            "metadata": body_extracted.get("metadata", {}) if isinstance(body_extracted.get("metadata"), dict) else {},
            "warnings": top_warnings,
            "email_context": email_content,
            "body_extracted": body_extracted,
        }

        if leave_lines_extracted is not None:
            extracted_json["leave_lines_extracted"] = leave_lines_extracted
            for rec in leave_lines_extracted.get("data", []):
                if isinstance(rec, dict):
                    rec_copy = dict(rec)
                    rec_copy["__extraction_type__"] = "leave_lines"
                    rec_copy["__source_file__"] = filename
                    extracted_json["data"].append(rec_copy)

        blocks.append(
            SourceBlock(order=order, type=BlockType.EXTRACTED_JSON, content=extracted_json, meta={})
        )
        order += 1

        logger.debug(
            "extract_email returning %d derived files: %s",
            len(self.derived_files),
            self.derived_files,
        )

        return SourceDoc(
            source_id=source_id,
            filename=filename,
            file_path=file_path,
            source_type="email",
            blocks=blocks,
            extracted=extracted_json,
        )

    # ------------------------------------------------------------------
    # 附件/内联资源导出子流程
    # ------------------------------------------------------------------

    # 与 CID 解析共享的临时状态：用于在主流程中回填 CID→文件路径映射与导出摘要
    _last_cid_to_path: Dict[str, str] = {}
    _last_attachments_summary: List[dict] = []
    _last_export_warnings: List[str] = []

    def _export_attachments(
        self,
        file_path: str,
        out_dir: str,
        blocks: List[SourceBlock],
        order: int,
        cid_list: List[str],
    ) -> int:
        """
        重新解析 `.eml` 并将所有 MIME parts 导出到 `out_dir`。

        该函数会追加导出文件对应的 `SourceBlock` 到传入的列表中，并维护以下状态：
        - `self._last_cid_to_path`：content-id（去掉尖括号）→ 文件路径
        - `self._last_attachments_summary`：附件摘要（仅 disposition=attachment）
        - `self._last_export_warnings`：导出失败/路径缺失等诊断信息

        返回：
        - 更新后的 `order`，用于保证块顺序连续。
        """
        msg = EmailParser.parse_raw(file_path)

        sha256_to_path: Dict[str, str] = {}
        cid_to_path: Dict[str, str] = {}
        derived_set: set = set()
        exported_sha_added: set = set()
        attachments_summary: List[dict] = []
        export_warnings: List[str] = []
        part_counter = {"index": 0}

        def _record_warning(part, reason: str) -> None:
            idx = part_counter.get("index", 0)
            export_warnings.append(
                f"export_failed idx={idx} filename={part.get_filename() or ''} "
                f"content_type={part.get_content_type()} reason={reason}"
            )

        def _iter_rfc822(part):
            payload = part.get_payload()
            if isinstance(payload, list):
                for item in payload:
                    if item:
                        yield item
            elif payload is not None:
                if hasattr(payload, "get_payload"):
                    yield payload
                elif isinstance(payload, (bytes, bytearray)):
                    try:
                        parser = BytesParser(policy=policy.default)
                        yield parser.parsebytes(payload)
                    except Exception:
                        pass

        def _process(part, cur_order: int) -> int:
            part_counter["index"] += 1
            content_type = part.get_content_type()

            if content_type == "message/rfc822":
                nested_found = False
                for nested_msg in _iter_rfc822(part):
                    nested_found = True
                    cur_order = _process(nested_msg, cur_order)
                if not nested_found:
                    _record_warning(part, "rfc822_payload_missing")

            exported = AttachmentHandler.export_eml_part(part, out_dir, sha256_to_path)
            if exported:
                is_img = AttachmentHandler.is_image_file(
                    exported.get("content_type", ""),
                    exported.get("filename", ""),
                )

                written_path = exported.get("written_path") or exported.get("path")
                if written_path and os.path.exists(written_path):
                    abs_path = os.path.abspath(written_path)
                    if abs_path not in derived_set:
                        derived_set.add(abs_path)
                        self.derived_files.append(abs_path)
                else:
                    _record_warning(part, "written_path_missing")

                if exported.get("content_id"):
                    cid_key = exported["content_id"].strip("<>")
                    cid_to_path[cid_key] = written_path or exported.get("path")

                if exported.get("disposition") == "attachment":
                    attachments_summary.append({
                        "filename": exported.get("filename", ""),
                        "content_type": exported.get("content_type", ""),
                        "size_bytes": exported.get("size_bytes", 0),
                    })

                sha256 = exported.get("sha256")
                if sha256 and sha256 in exported_sha_added:
                    # 内容去重：同一二进制内容可能在不同部件中重复出现，避免重复写入块
                    pass
                else:
                    if sha256:
                        exported_sha_added.add(sha256)
                    block_type = BlockType.EML_IMAGE_FILE if is_img else BlockType.EML_FILE_PART
                    blocks.append(
                        SourceBlock(order=cur_order, type=block_type, content=exported, meta={})
                    )
                    cur_order += 1

            if part.is_multipart():
                for subpart in part.iter_parts():
                    cur_order = _process(subpart, cur_order)

            return cur_order

        order = _process(msg, order)

        # 保存导出结果，供主流程后续生成摘要与 CID→文件路径映射
        self._last_cid_to_path = cid_to_path
        self._last_attachments_summary = attachments_summary
        self._last_export_warnings = export_warnings

        logger.info(
            "EML exported parts: (derived_files=%d attachments=%d)",
            len(self.derived_files),
            len(attachments_summary),
        )
        return order

    # ------------------------------------------------------------------
    # CID 解析
    # ------------------------------------------------------------------

    @staticmethod
    def _resolve_cid(cid: str, cid_to_path: Dict[str, str]) -> Optional[str]:
        """
        将 CID 解析为导出文件路径（带大小写不敏感兜底）。

        说明：
        - Content-ID 在邮件中可能带尖括号 `<...>`；
        - 实际存储的 key 与 HTML 引用的 CID 可能大小写不同；
        - 这里做规范化与遍历匹配，尽可能提高命中率。
        """
        cid_normalized = cid.strip("<>")
        if cid in cid_to_path:
            return cid_to_path[cid]
        if cid_normalized in cid_to_path:
            return cid_to_path[cid_normalized]
        for stored_cid, path in cid_to_path.items():
            if stored_cid.lower() == cid.lower() or stored_cid.lower() == cid_normalized.lower():
                return path
        return None

    # ------------------------------------------------------------------
    # LLM 正文抽取
    # ------------------------------------------------------------------

    def _llm_extract_body(
        self,
        normalized_body_text: str,
        body_text_meta: dict,
        filename: str,
    ) -> dict:
        """
        调用 LLM 从邮件正文中抽取通用结构化记录。

        输入：
        - normalized_body_text：归一化后的正文纯文本（已做必要清洗/截断）
        - body_text_meta：正文清洗产生的元信息（用于回填到抽取结果中，便于追踪与调试）
        - filename：源文件名（会写入每条记录的 `__source_file__` 字段）

        输出（形状保证）：
        - 通过 `ContentCleaner.coerce_body_extracted_shape()` 对 LLM 原始输出进行整形，
          以保证返回结果至少包含：`data`（列表）、`metadata`（字典）、`warnings`（列表）。

        容错策略：
        - 正文为空或缺少提示词模板（`EML_BODY_TO_JSON_PROMPT`）时，不调用 LLM，返回空数据并写入告警（`warnings`）；
        - LLM 调用异常会捕获并记录到告警（`warnings`），避免中断整个邮件抽取流程。
        """
        body_warnings: List[str] = []
        prompt_tpl = (self.prompts or {}).get("EML_BODY_TO_JSON_PROMPT", "")
        body_llm_raw: Any = None

        if not normalized_body_text.strip():
            body_warnings.append("body_extracted: empty normalized body text")
        elif not prompt_tpl.strip():
            body_warnings.append("body_extracted: missing prompt EML_BODY_TO_JSON_PROMPT")
        else:
            prompt = prompt_tpl + "\n\nEMAIL_BODY_TEXT:\n" + normalized_body_text
            try:
                logger.info("Email body extraction: LLM start (%s)", filename)
                if hasattr(self.llm, "chat_json_once"):
                    body_llm_raw = self.llm.chat_json_once(
                        prompt,
                        system=None,
                        temperature=0,
                        timeout=30.0,
                        step="email_body_to_json",
                    )
                else:
                    body_llm_raw = self.llm.chat_json(
                        prompt,
                        system=None,
                        temperature=0,
                        step="email_body_to_json",
                    )
                logger.info("Email body extraction: LLM done (%s)", filename)
            except Exception as e:
                body_warnings.append(f"body_extracted: llm_exception: {type(e).__name__}: {str(e)[:200]}")
                body_llm_raw = {"error": "llm_exception", "message": str(e)[:200]}

        body_extracted = ContentCleaner.coerce_body_extracted_shape(body_llm_raw, body_text_meta, body_warnings)

        normalized_records: List[dict] = []
        for item in body_extracted.get("data", []) if isinstance(body_extracted, dict) else []:
            if isinstance(item, dict):
                rec = ContentCleaner.stringify_record_values(item)
                rec["__source_file__"] = filename
                normalized_records.append(rec)
        body_extracted["data"] = normalized_records

        return body_extracted

    # ------------------------------------------------------------------
    # `.txt` 抽取流程
    # ------------------------------------------------------------------

    def _extract_txt(self, file_path: str, filename: str, source_id: str) -> SourceDoc:
        """
        抽取 `.txt` 邮件文本文件内容。

        处理方式：
        - 读取全文作为 `BlockType.TEXT` block；
        - 使用 `EMAIL_TO_JSON_PROMPT` 调用 LLM，将文本抽取为结构化 JSON；
        - 将 LLM 输出作为 `BlockType.EXTRACTED_JSON` block 写入 `SourceDoc`。
        """
        blocks: List[SourceBlock] = []

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        blocks.append(SourceBlock(order=1, type=BlockType.TEXT, content=text, meta={}))

        prompt = self.prompts["EMAIL_TO_JSON_PROMPT"] + "\n\nEMAIL_TEXT:\n" + text
        extracted_json = self.llm.chat_json(prompt, system=None, step="email_to_json")

        blocks.append(SourceBlock(order=2, type=BlockType.EXTRACTED_JSON, content=extracted_json, meta={}))

        return SourceDoc(
            source_id=source_id,
            filename=filename,
            file_path=file_path,
            source_type="email",
            blocks=blocks,
            extracted=extracted_json,
        )

    # ------------------------------------------------------------------
    # `.docx` 抽取流程
    # ------------------------------------------------------------------

    def _extract_docx(self, file_path: str, filename: str, source_id: str) -> SourceDoc:
        """
        抽取 `.docx` 文件内容。

        处理方式：
        - 段落：提取非空段落文本；
        - 表格：将首行视为表头，后续行按表头映射为字典（表头不足时用 `col_i` 兜底）；
        - 图片：仅生成图片占位信息（不导出二进制），用于让下游感知文档包含图片资源。

        最终会将上述结构序列化为 `DOCX_BLOCKS_JSON`，再使用 `EMAIL_TO_JSON_PROMPT` 调用 LLM 抽取结构化 JSON。
        """
        blocks: List[SourceBlock] = []

        doc = Document(file_path)
        docx_blocks: List[dict] = []

        for para in doc.paragraphs:
            if para.text.strip():
                docx_blocks.append({"type": "text", "content": para.text.strip()})

        for table in doc.tables:
            if len(table.rows) > 0:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                rows = []
                for row in table.rows[1:]:
                    row_data = {}
                    for i, cell in enumerate(row.cells):
                        key = headers[i] if i < len(headers) else f"col_{i}"
                        row_data[key] = cell.text.strip()
                    rows.append(row_data)
                docx_blocks.append({"type": "table", "content": rows})

        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                docx_blocks.append({
                    "type": "image_placeholder",
                    "content": {"note": "docx image detected", "index": image_count},
                })
                image_count += 1

        docx_type_map = {
            "text": BlockType.TEXT,
            "table": BlockType.TABLE,
            "image_placeholder": BlockType.IMAGE_PLACEHOLDER,
        }
        for i, block_data in enumerate(docx_blocks, start=1):
            block_type = docx_type_map.get(block_data["type"], BlockType.TEXT)
            blocks.append(SourceBlock(order=i, type=block_type, content=block_data["content"], meta={}))

        blocks_json = json.dumps(docx_blocks, ensure_ascii=False, separators=(",", ":"))

        prompt = self.prompts["EMAIL_TO_JSON_PROMPT"] + "\n\nDOCX_BLOCKS_JSON:\n" + blocks_json
        extracted_json = self.llm.chat_json(prompt, system=None, step="email_to_json")

        blocks.append(
            SourceBlock(order=len(blocks) + 1, type=BlockType.EXTRACTED_JSON, content=extracted_json, meta={})
        )

        return SourceDoc(
            source_id=source_id,
            filename=filename,
            file_path=file_path,
            source_type="email",
            blocks=blocks,
            extracted=extracted_json,
        )
